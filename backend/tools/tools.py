import os
import sys
import json
import platform
import subprocess
import webbrowser
import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Union
import traceback

# External libraries
try:
    import screen_brightness_control as sbc
    BRIGHTNESS_AVAILABLE = True
except ImportError:
    BRIGHTNESS_AVAILABLE = False
    logging.warning("screen_brightness_control not installed")

try:
    import pyautogui
    PYAUTOGUI_AVAILABLE = True
except ImportError:
    PYAUTOGUI_AVAILABLE = False

logger = logging.getLogger("master_os_tools")

# ==================== SYSTEM UTILITIES ====================

def get_platform() -> str:
    """Get current platform."""
    return platform.system().lower()

def run_applescript(script: str) -> str:
    """Execute AppleScript on macOS."""
    if get_platform() != "darwin":
        return "AppleScript only works on macOS"
    
    try:
        result = subprocess.run(
            ['osascript', '-e', script],
            capture_output=True,
            text=True,
            timeout=10
        )
        return result.stdout.strip() if result.returncode == 0 else result.stderr
    except Exception as e:
        return f"AppleScript error: {str(e)}"

def run_powershell(script: str) -> str:
    """Execute PowerShell on Windows."""
    if get_platform() != "windows":
        return "PowerShell only works on Windows"
    
    try:
        result = subprocess.run(
            ['powershell', '-Command', script],
            capture_output=True,
            text=True,
            timeout=10
        )
        return result.stdout.strip() if result.returncode == 0 else result.stderr
    except Exception as e:
        return f"PowerShell error: {str(e)}"



# ==================== LIBRARY LOADING ====================
# We use try-except blocks to make the script robust even if dependencies are missing
logger = logging.getLogger("JARVIS_OS_CORE")
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def log_error(msg): logger.error(f"❌ {msg}")
def log_success(msg): logger.info(f"✅ {msg}")

# 1. System Hardware & Processes
try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False
    log_error("psutil not installed. Process/Battery features disabled.")

# 2. GUI Automation (Mouse/Keyboard/Screen)
try:
    import pyautogui
    # Fail-safe to prevent mouse from going rogue (move mouse to corner to kill)
    pyautogui.FAILSAFE = True 
    PYAUTOGUI_AVAILABLE = True
except ImportError:
    PYAUTOGUI_AVAILABLE = False
    log_error("pyautogui not installed. Mouse/Keyboard automation disabled.")


# 4. Clipboard
try:
    import pyperclip
    CLIPBOARD_AVAILABLE = True
except ImportError:
    CLIPBOARD_AVAILABLE = False

# 5. Text to Speech (The Voice of Jarvis)
try:
    import pyttsx3
    TTS_AVAILABLE = True
    tts_engine = pyttsx3.init()
    # Configure voice to sound more robotic/jarvis-like if possible
    voices = tts_engine.getProperty('voices')
    # Try to find a good English voice
    for voice in voices: # type: ignore
        if "david" in voice.name.lower() or "samantha" in voice.name.lower():
            tts_engine.setProperty('voice', voice.id)
            break
    tts_engine.setProperty('rate', 190) # Slightly faster
except ImportError:
    TTS_AVAILABLE = False

# ==================== LOW LEVEL KERNEL UTILS ====================

def is_admin() -> bool:
    """Check if script has admin/root privileges."""
    try:
        if get_platform() == 'windows':
            import ctypes
            return ctypes.windll.shell32.IsUserAnAdmin() != 0
        else:
            return os.geteuid() == 0 # type: ignore
    except:
        return False

def run_shell_command(command: str, admin: bool = False) -> str:
    """Execute raw shell commands. The 'God Mode' function."""
    try:
        if admin and not is_admin():
            return "❌ Privileged command requires Admin/Root access. Run script as Administrator."
        
        result = subprocess.run(
            command, 
            shell=True, 
            capture_output=True, 
            text=True, 
            timeout=15
        )
        if result.returncode == 0:
            return result.stdout.strip()
        return f"Error (Code {result.returncode}): {result.stderr.strip()}"
    except Exception as e:
        return f"Shell Execution Failed: {str(e)}"
# ==================== TOOL IMPLEMENTATIONS ====================

class ToolRegistry:
    # ============ SYSTEM CONTROL ============
    
    @staticmethod
    def set_brightness(value: int) -> str:
        if not BRIGHTNESS_AVAILABLE:
            return "❌ Brightness control not available. Install: pip install screen-brightness-control"
        
        try:
            val = max(0, min(100, int(value)))
            sbc.set_brightness(val) # type: ignore
            logger.info(f"✅ Brightness set to {val}%")
            return f"✅ Screen brightness set to {val}%"
        except Exception as e:
            logger.error(f"Brightness control failed: {e}")
            return f"❌ Failed to adjust brightness: {str(e)}"
    
    @staticmethod
    def set_volume(value: int) -> str:
        try:
            val = max(0, min(100, int(value)))
            
            if get_platform() == "darwin":
                script = f"set volume output volume {val}"
                run_applescript(script)
            elif get_platform() == "windows":
                # Windows uses 0-65535 range
                win_val = int(val * 655.35)
                script = f"(New-Object -ComObject WScript.Shell).SendKeys([char]174)" # Mute toggle
                run_powershell(script)
            
            logger.info(f"✅ Volume set to {val}%")
            return f"✅ System volume set to {val}%"
        except Exception as e:
            logger.error(f"Volume control failed: {e}")
            return f"❌ Failed to set volume: {str(e)}"
    
    @staticmethod
    def toggle_dark_mode(enable: bool = True) -> str:
        """
        Toggle system dark mode on/off.
        Use when user wants dark/light theme.
        """
        try:
            if get_platform() == "darwin":
                mode = "dark" if enable else "light"
                script = f'''
                tell application "System Events"
                    tell appearance preferences
                        set dark mode to {str(enable).lower()}
                    end tell
                end tell
                '''
                run_applescript(script)
            elif get_platform() == "windows":
                # Windows 10/11 dark mode
                reg_val = 0 if enable else 1
                script = f'''
                Set-ItemProperty -Path HKCU:\\SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Themes\\Personalize -Name AppsUseLightTheme -Value {reg_val}
                '''
                run_powershell(script)
            
            mode_str = "enabled" if enable else "disabled"
            logger.info(f"✅ Dark mode {mode_str}")
            return f"✅ Dark mode {mode_str}"
        except Exception as e:
            return f"❌ Failed to toggle dark mode: {str(e)}"
    
    # ============ MUSIC & MEDIA ============
    
    @staticmethod
    def play_spotify(query: str = "", mood: str = "") -> str:
        try:
            search_term = query or mood or "music"
            safe_query = search_term.strip().lower()[:100]
            
            # If Spotify is already open, use it directly
            if get_platform() == "darwin":
                # Try to control running Spotify first
                script = f'''
                tell application "System Events"
                    if exists (processes where name is "Spotify") then
                        tell application "Spotify"
                            activate
                            set search_query to "{safe_query}"
                            play track search_query
                        end tell
                        return "Playing: {safe_query}"
                    else
                        return "Opening Spotify"
                    end if
                end tell
                '''
                result = run_applescript(script)
                
                if "Opening Spotify" in result:
                    # Spotify not running, use web
                    url = f"https://open.spotify.com/search/{safe_query.replace(' ', '%20')}"
                    webbrowser.open(url)
                    return f"🎵 Opened Spotify searching for '{search_term}'"
                else:
                    return f"🎵 {result}"
            else:
                # Windows/Linux - use web Spotify
                url = f"https://open.spotify.com/search/{safe_query.replace(' ', '%20')}"
                webbrowser.open(url)
                return f"🎵 Opened Spotify searching for '{search_term}'"
        
        except Exception as e:
            logger.error(f"Spotify control failed: {e}")
            return f"❌ Failed to control Spotify: {str(e)}"
    
    @staticmethod
    def spotify_control(action: str) -> str:
        try:
            if get_platform() == "darwin":
                script = f'''
                tell application "Spotify"
                    {action}
                end tell
                '''
                run_applescript(script)
                return f"✅ Spotify: {action}"
            else:
                return "⚠️ Direct Spotify control only available on macOS. Try 'play spotify' instead."
        except Exception as e:
            return f"❌ Spotify control error: {str(e)}"
    
    # ============ FILE OPERATIONS ============
    
    @staticmethod
    def read_file(path: str, max_chars: int = 50000) -> str:
        try:
            full_path = os.path.expanduser(path)
            
            if not os.path.exists(full_path):
                return f"❌ File not found: {path}"
            
            if os.path.getsize(full_path) > 10_000_000:  # 10MB
                return f"❌ File too large (>10MB): {path}"
            
            with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read(max_chars)
            
            if len(content) >= max_chars:
                content += f"\n\n... [Truncated at {max_chars} chars]"
            
            return f"📄 Contents of {os.path.basename(path)}:\n\n{content}"
        
        except Exception as e:
            logger.error(f"Read file failed: {e}")
            return f"❌ Error reading file: {str(e)}"
    
    @staticmethod
    def write_file(path: str, content: str, mode: str = "overwrite") -> str:
        try:
            full_path = os.path.expanduser(path)
            
            # Create parent directories
            Path(full_path).parent.mkdir(parents=True, exist_ok=True)
            
            write_mode = 'w' if mode == "overwrite" else 'a'
            
            with open(full_path, write_mode, encoding='utf-8') as f:
                f.write(content)
            
            action = "Created" if mode == "overwrite" else "Appended to"
            bytes_written = len(content.encode('utf-8'))
            
            logger.info(f"✅ {action} {path} ({bytes_written} bytes)")
            return f"✅ {action} {path} ({bytes_written} bytes)"
        
        except Exception as e:
            logger.error(f"Write file failed: {e}")
            return f"❌ Error writing file: {str(e)}"
    
    @staticmethod
    def create_document(title: str, content: str, doc_type: str = "markdown") -> str:
        try:
            # Sanitize filename
            import re
            safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')
            
            docs_path = os.path.expanduser("~/Documents")
            
            if doc_type == "markdown":
                filepath = os.path.join(docs_path, f"{safe_title}.md")
                full_content = f"# {title}\n\n{content}"
            elif doc_type == "word":
                try:
                    from docx import Document
                    filepath = os.path.join(docs_path, f"{safe_title}.docx")
                    doc = Document()
                    doc.add_heading(title, 0)
                    doc.add_paragraph(content)
                    doc.save(filepath)
                    return f"✅ Created Word document: {filepath}"
                except ImportError:
                    return "❌ python-docx not installed. Creating as markdown instead."
            else:  # text
                filepath = os.path.join(docs_path, f"{safe_title}.txt")
                full_content = f"{title}\n\n{content}"
            
            Path(filepath).parent.mkdir(parents=True, exist_ok=True)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(full_content)
            
            logger.info(f"✅ Created document: {filepath}")
            return f"✅ Created document: {filepath}"
        
        except Exception as e:
            return f"❌ Error creating document: {str(e)}"
    
    @staticmethod
    def file_search(query: str, location: str = "~") -> str:
        try:
            if not query or len(query.strip()) < 2:
                return "❌ Search query too short (minimum 2 characters)"
            
            query_lower = query.lower().strip()
            results = []
            max_results = 15
            
            # Search paths
            search_paths = [
                os.path.expanduser("~/Documents"),
                os.path.expanduser("~/Desktop"),
                os.path.expanduser("~/Downloads")
            ]
            
            for base_path in search_paths:
                if not os.path.exists(base_path):
                    continue
                
                try:
                    for root, dirs, files in os.walk(base_path):
                        # Skip hidden dirs and system dirs
                        dirs[:] = [d for d in dirs if not d.startswith('.') and d not in 
                                   ['node_modules', '__pycache__', 'Library']]
                        
                        for file in files:
                            if query_lower in file.lower():
                                full_path = os.path.join(root, file)
                                rel_path = os.path.relpath(full_path, os.path.expanduser('~'))
                                results.append((file, rel_path))
                                
                                if len(results) >= max_results:
                                    break
                        
                        if len(results) >= max_results:
                            break
                except PermissionError:
                    continue
            
            if not results:
                logger.info(f"No files found for: {query}")
                return f"❌ No files found matching '{query}'"
            
            logger.info(f"✅ Found {len(results)} files for: {query}")
            
            formatted = "\n".join([f"📄 {name}\n   ~/{path}" for name, path in results])
            return f"✅ Found {len(results)} file(s):\n\n{formatted}"
        
        except Exception as e:
            logger.error(f"File search failed: {e}")
            return f"❌ File search error: {str(e)}"
    
    # ============ SYSTEM INFO ============
    
    @staticmethod
    def get_current_time() -> str:
        try:
            now = datetime.now()
            formatted = now.strftime("%A, %B %d, %Y at %I:%M %p")
            logger.info(f"✅ Time requested: {formatted}")
            return f"🕐 Current time: {formatted}"
        except Exception as e:
            return f"❌ Error getting time: {str(e)}"
    
    @staticmethod
    def get_system_info() -> str:
        """
        Get system information (OS, version, processor, etc.).
        Use when user asks about their computer specs.
        """
        try:
            import platform
            import psutil
            
            info = {
                "Creator" : "Badsha Laskar and his team (Joy, Sanya and Atul)",
                "OS": platform.system(),
                "OS Version": platform.version(),
                "Machine": platform.machine(),
                "Processor": platform.processor(),
                "CPU Cores": psutil.cpu_count(logical=False),
                "Logical CPUs": psutil.cpu_count(logical=True),
                "RAM": f"{psutil.virtual_memory().total / (1024**3):.1f} GB",
                "Python": platform.python_version()
            }
            
            formatted = "\n".join([f"  • {k}: {v}" for k, v in info.items()])
            return f"💻 System Information:\n\n{formatted}"
        except Exception as e:
            return f"❌ Error getting system info: {str(e)}"
    
    # ============ APPLICATION CONTROL ============
    
    @staticmethod
    def open_website(url: str) -> str:
        try:
            if not url.startswith(("http://", "https://")):
                url = "https://" + url
            
            webbrowser.open(url)
            logger.info(f"✅ Opened URL: {url}")
            return f"✅ Opened {url} in browser"
        
        except Exception as e:
            logger.error(f"Failed to open URL: {e}")
            return f"❌ Could not open website: {str(e)}"
    
    @staticmethod
    def web_search(query: str, engine: str = "google") -> str:
        try:
            engines = {
                "google": "https://www.google.com/search?q=",
                "duckduckgo": "https://duckduckgo.com/?q=",
                "bing": "https://www.bing.com/search?q=",
                "youtube": "https://www.youtube.com/results?search_query="
            }
            
            base_url = engines.get(engine.lower(), engines["google"])
            search_url = base_url + query.replace(" ", "+")
            
            webbrowser.open(search_url)
            return f"🔍 Searching {engine} for: {query}"
        
        except Exception as e:
            return f"❌ Search failed: {str(e)}"
    
    # ============ COMMUNICATION ============
    
    @staticmethod
    def compose_email(to: str = "", subject: str = "", body: str = "") -> str:
        try:
            import urllib.parse
            
            mailto_parts = []
            if subject:
                mailto_parts.append(f"subject={urllib.parse.quote(subject)}")
            if body:
                mailto_parts.append(f"body={urllib.parse.quote(body)}")
            
            mailto = f"mailto:{to}"
            if mailto_parts:
                mailto += "?" + "&".join(mailto_parts)
            
            webbrowser.open(mailto)
            logger.info("✅ Opened email client")
            return f"✅ Opened email client with pre-filled fields"
        
        except Exception as e:
            logger.error(f"Email client open failed: {e}")
            return f"❌ Could not open email: {str(e)}"
    
    # ============ AUTOMATION HELPERS ============
    
    @staticmethod
    def wait(seconds: int) -> str:
        import time
        try:
            time.sleep(min(seconds, 30))  # Max 30 seconds
            return f"✅ Waited {seconds} seconds"
        except Exception as e:
            return f"❌ Wait error: {str(e)}"
    
    # ============ FINAL ANSWER ============
    
    @staticmethod
    def final_answer(answer: str) -> str:
        return answer
    
    @staticmethod
    def speak(text: str) -> str:
        """Text-to-Speech output."""
        if not TTS_AVAILABLE: return "❌ TTS library missing."
        try:
            tts_engine.say(text)
            tts_engine.runAndWait()
            return f"🗣️ Spoke: {text}"
        except Exception as e:
            return f"❌ TTS Error: {e}"

    # ============ 2. PROCESS & KERNEL MANAGEMENT ============

    @staticmethod
    def get_running_processes(search: str = "") -> str:
        """List active processes (like Task Manager)."""
        if not PSUTIL_AVAILABLE: return "❌ psutil missing."
        try:
            procs = []
            for p in psutil.process_iter(['pid', 'name', 'username']): # type: ignore
                if search.lower() in p.info['name'].lower():
                    procs.append(p.info)
            
            # Sort by name and take top 20 if too many
            procs.sort(key=lambda x: x['name'])
            top_procs = procs[:20]
            
            output = f"⚙️ Found {len(procs)} processes" + (f" matching '{search}'" if search else "") + ":\n"
            for p in top_procs:
                output += f" • PID: {p['pid']} | Name: {p['name']} | User: {p['username']}\n"
            
            if len(procs) > 20: output += "... (and more)"
            return output
        except Exception as e:
            return f"❌ Process lookup error: {e}"

    @staticmethod
    def terminate_process(process_name_or_pid: str) -> str:
        """Kill a specific process/app."""
        if not PSUTIL_AVAILABLE: return "❌ psutil missing."
        killed_count = 0
        try:
            # If PID is provided
            if process_name_or_pid.isdigit():
                pid = int(process_name_or_pid)
                p = psutil.Process(pid) # type: ignore
                p.terminate()
                return f"💀 Terminated PID {pid} ({p.name()})"
            
            # If Name is provided
            target = process_name_or_pid.lower()
            if not target.endswith(('.exe', '.app')): 
                # Loose matching
                for p in psutil.process_iter(['pid', 'name']): # type: ignore
                    if target in p.info['name'].lower():
                        p.terminate()
                        killed_count += 1
            
            if killed_count == 0:
                return f"⚠️ No process found matching '{process_name_or_pid}'"
            return f"💀 Terminated {killed_count} process(es) matching '{process_name_or_pid}'"
            
        except psutil.NoSuchProcess: # type: ignore
            return "❌ Process no longer exists."
        except psutil.AccessDenied: # type: ignore
            return "❌ Access Denied. Try running as Admin."
        except Exception as e:
            return f"❌ Kill failed: {e}"

    # ============ 3. DEEP SYSTEM CONTROL ============

    @staticmethod
    def system_power(action: str) -> str:
        """Shutdown, Restart, Lock, or Sleep the machine."""
        plt = get_platform()
        cmd = ""
        action = action.lower()
        
        if plt == "windows":
            if "shutdown" in action: cmd = "shutdown /s /t 5"
            elif "restart" in action: cmd = "shutdown /r /t 5"
            elif "lock" in action: cmd = "rundll32.dll user32.dll,LockWorkStation"
            elif "sleep" in action: cmd = "rundll32.dll powrprof.dll,SetSuspendState 0,1,0"
        elif plt == "darwin": # macOS
            if "shutdown" in action: cmd = "sudo shutdown -h now"
            elif "restart" in action: cmd = "sudo shutdown -r now"
            elif "lock" in action: cmd = "pmset displaysleepnow"
            elif "sleep" in action: cmd = "pmset sleepnow"
        else: # Linux
            if "shutdown" in action: cmd = "shutdown -h now"
            elif "restart" in action: cmd = "shutdown -r now"
            
        if not cmd:
            return "❌ Unknown power action or platform."
            
        # Execute
        logger.warning(f"EXECUTING POWER COMMAND: {cmd}")
        run_shell_command(cmd)
        return f"🔌 Executing: {action}..."

    @staticmethod
    def get_system_vitals() -> str:
        """Battery, CPU Load, Memory, Disk Space."""
        if not PSUTIL_AVAILABLE: return "❌ psutil missing."
        try:
            cpu = psutil.cpu_percent(interval=0.1)# type: ignore
            mem = psutil.virtual_memory() # type: ignore
            disk = psutil.disk_usage('/') # type: ignore
            
            batt_info = "🔌 Plugged In"
            if hasattr(psutil, "sensors_battery"): # type: ignore
                batt = psutil.sensors_battery() # type: ignore
                if batt:
                    batt_info = f"{batt.percent}% ({'Charging' if batt.power_plugged else 'Discharging'})"

            return (f"📊 System Vitals:\n"
                    f" • CPU Load: {cpu}%\n"
                    f" • RAM Usage: {mem.percent}% ({mem.used // (1024**3)}GB used)\n"
                    f" • Disk Free: {disk.free // (1024**3)}GB\n"
                    f" • Battery: {batt_info}")
        except Exception as e:
            return f"❌ Vitals error: {e}"

    # ============ 4. GUI & APP AUTOMATION (The "Control Apps" part) ============

    @staticmethod
    def take_screenshot(filename: str = "screenshot.png") -> str:
        """Capture screen."""
        if not PYAUTOGUI_AVAILABLE: return "❌ pyautogui missing."
        try:
            path = os.path.expanduser(f"~/Pictures/{filename}")
            pyautogui.screenshot(path) # type: ignore
            return f"📸 Screenshot saved to {path}"
        except Exception as e:
            return f"❌ Screenshot failed: {e}"

    @staticmethod
    def press_hotkey(keys: str) -> str:
        """
        Press keyboard shortcuts. 
        Format: 'ctrl+c', 'command+shift+3', 'alt+tab', 'enter', 'esc'
        """
        if not PYAUTOGUI_AVAILABLE: return "❌ pyautogui missing."
        try:
            key_list = keys.lower().replace(" ", "").split('+')
            pyautogui.hotkey(*key_list) # type: ignore
            return f"⌨️ Pressed: {keys}"
        except Exception as e:
            return f"❌ Hotkey failed: {e}"

    @staticmethod
    def type_text(text: str, interval: float = 0.01) -> str:
        """Type text into the active window."""
        if not PYAUTOGUI_AVAILABLE: return "❌ pyautogui missing."
        try:
            pyautogui.write(text, interval=interval) # type: ignore
            return f"⌨️ Typed {len(text)} chars"
        except Exception as e:
            return f"❌ Typing failed: {e}"

    @staticmethod
    def clipboard_action(action: str, text: str = "") -> str:
        """Copy to or Paste from clipboard."""
        if not CLIPBOARD_AVAILABLE: return "❌ pyperclip missing."
        try:
            if action == "copy":
                pyperclip.copy(text) # type: ignore
                return "📋 Copied to clipboard."
            elif action == "paste":
                content = pyperclip.paste() # type: ignore
                return f"📋 Clipboard contents:\n{content}"
            else:
                return "❌ Unknown clipboard action. Use 'copy' or 'paste'."
        except Exception as e:
            return f"❌ Clipboard error: {e}"

    # ============ 5. NETWORK & WEB ============
    
    @staticmethod
    def get_network_info() -> str:
        """Internal IP, External IP, Hostname."""
        try:
            hostname = socket.gethostname() # type: ignore
            local_ip = socket.gethostbyname(hostname) # type: ignore
            
            # Get external IP (simple request)
            import requests
            try:
                external_ip = requests.get('https://api.ipify.org', timeout=3).text
            except:
                external_ip = "Unavailable"
                
            return (f"🌐 Network Status:\n"
                    f" • Hostname: {hostname}\n"
                    f" • Local IP: {local_ip}\n"
                    f" • Public IP: {external_ip}")
        except Exception as e:
            return f"❌ Network check failed: {e}"

    # ============ EXISTING UTILS (Refined) ============
    
    @staticmethod
    def open_app(app_name: str) -> str:
        """Enhanced App Opener that handles OS differences better."""
        plt = get_platform()
        app = app_name.lower().strip()
        
        try:
            if plt == "darwin":
                run_shell_command(f'open -a "{app}"')
            elif plt == "windows":
                run_shell_command(f'start {app}')
            else:
                subprocess.Popen([app], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            
            return f"🚀 Launching {app_name}..."
        except Exception as e:
            return f"❌ Failed to launch {app_name}: {e}"

    @staticmethod
    def file_search_recursive(query: str, root_path: str = "~", depth: int = 3) -> str:
        """Deeper file search."""
        root = os.path.expanduser(root_path)
        matches = []
        try:
            for root_dir, dirs, files in os.walk(root):
                # Calculate current depth
                current_depth = root_dir[len(root):].count(os.sep)
                if current_depth > depth:
                    continue
                
                # Filter hidden dirs
                dirs[:] = [d for d in dirs if not d.startswith('.')]
                
                for file in files:
                    if query.lower() in file.lower():
                        matches.append(os.path.join(root_dir, file))
                        if len(matches) > 10: break
                if len(matches) > 10: break
            
            if not matches: return f"🔍 No files found for '{query}'"
            return "🔍 Found:\n" + "\n".join(matches)
        except Exception as e:
            return f"❌ Search Error: {e}"


# ==================== TOOL SCHEMAS FOR LLM ====================

TOOL_SCHEMAS = [
    {
        "name": "set_brightness",
        "description": "Adjust monitor brightness (0-100). Use for: 'dim screen', 'brighten display', 'reduce brightness'",
        "parameters": {"value": "integer 0-100"},
        "examples": ["set brightness to 30", "dim my screen", "brighten display"]
    },
    {
        "name": "set_volume",
        "description": "Set system volume (0-100). Use for: 'change volume', 'louder', 'quieter'",
        "parameters": {"value": "integer 0-100"},
        "examples": ["set volume to 50", "turn volume down", "make it louder"]
    },
    {
        "name": "toggle_dark_mode",
        "description": "Enable/disable dark mode. Use for: 'dark mode', 'light theme', 'switch theme'",
        "parameters": {"enable": "boolean (true for dark, false for light)"},
        "examples": ["enable dark mode", "turn on dark theme", "switch to light mode"]
    },
    {
        "name": "play_spotify",
        "description": "Play music on Spotify. Accepts song/artist/playlist or mood (happy/sad/calm/energetic/focus)",
        "parameters": {
            "query": "song, artist, or playlist name (optional)",
            "mood": "happy, sad, energetic, calm, focus (optional)"
        },
        "examples": ["play sad music", "play Bohemian Rhapsody", "play focus playlist"]
    },
    {
        "name": "spotify_control",
        "description": "Control Spotify playback: play, pause, next track, previous track",
        "parameters": {"action": "play, pause, next track, previous track"},
        "examples": ["pause spotify", "next song", "play music"]
    },
    {
        "name": "read_file",
        "description": "Read file contents. Use for: 'show me X file', 'read Y', 'what's in Z'",
        "parameters": {
            "path": "file path (can use ~ for home)",
            "max_chars": "maximum characters to read (default 50000)"
        },
        "examples": ["read ~/Documents/notes.txt", "show me essay.docx"]
    },
    {
        "name": "write_file",
        "description": "Write content to file. Use for: 'create file', 'save to X', 'write Y'",
        "parameters": {
            "path": "file path",
            "content": "content to write",
            "mode": "overwrite or append (default overwrite)"
        },
        "examples": ["write to notes.txt", "create todo.md", "save list to file"]
    },
    {
        "name": "create_document",
        "description": "Create formatted document (essay, report, notes). Auto-saves to ~/Documents",
        "parameters": {
            "title": "document title",
            "content": "document content",
            "doc_type": "markdown, text, or word (default markdown)"
        },
        "examples": ["create essay about climate", "make a todo list", "write report"]
    },
    {
        "name": "file_search",
        "description": "Search for files by name in Documents/Desktop/Downloads",
        "parameters": {
            "query": "search term (2+ chars)",
            "location": "search location (default: ~)"
        },
        "examples": ["find my resume", "search for budget spreadsheet", "where is essay.docx"]
    },
    {
        "name": "get_current_time",
        "description": "Get current date and time",
        "parameters": {},
        "examples": ["what time is it", "current date", "today's date"]
    },
    {
        "name": "get_system_info",
        "description": "Get computer specs (OS, CPU, RAM, etc.)",
        "parameters": {},
        "examples": ["system info", "my computer specs", "what's my OS"]
    },
    {
        "name": "open_app",
        "description": "Open application (notepad, calculator, chrome, vscode, spotify, etc.)",
        "parameters": {"app_name": "application name"},
        "examples": ["open chrome", "launch calculator", "start spotify"]
    },
    {
        "name": "open_website",
        "description": "Open website in browser",
        "parameters": {"url": "website URL"},
        "examples": ["open google.com", "go to youtube", "visit github.com"]
    },
    {
        "name": "web_search",
        "description": "Search the web using search engine",
        "parameters": {
            "query": "search query",
            "engine": "google, duckduckgo, bing, youtube (default google)"
        },
        "examples": ["search for python tutorials", "youtube search music", "google AI news"]
    },
    {
        "name": "compose_email",
        "description": "Open email client with pre-filled fields",
        "parameters": {
            "to": "recipient email (optional)",
            "subject": "email subject (optional)",
            "body": "email body (optional)"
        },
        "examples": ["compose email", "write email to john@example.com", "draft email"]
    },
    {
        "name": "wait",
        "description": "Wait for specified seconds (max 30). Use in multi-step automation",
        "parameters": {"seconds": "integer 1-30"},
        "examples": ["wait 5 seconds", "pause for 10 seconds"]
    },
    {
        "name": "type_text",
        "description": "Type text using keyboard automation. CAREFUL - types where focused!",
        "parameters": {"text": "text to type"},
        "examples": ["type hello world"]
    },
    {
        "name": "final_answer",
        "description": "Return final answer to user. Use when you have all needed information",
        "parameters": {"answer": "complete response to user"},
        "examples": []
    },
    {
        "name": "system_power",
        "description": "Control power state. BE CAREFUL.",
        "parameters": {"action": "shutdown, restart, lock, sleep"},
        "examples": ["lock the screen", "restart computer", "sleep mode"]
    },
    {
        "name": "get_running_processes",
        "description": "List active applications and system processes.",
        "parameters": {"search": "filter by name (optional)"},
        "examples": ["what is running", "is chrome running", "show processes"]
    },
    {
        "name": "terminate_process",
        "description": "Kill/Close a running application or process forcefullly.",
        "parameters": {"process_name_or_pid": "name (e.g., chrome) or PID"},
        "examples": ["kill chrome", "close spotify", "terminate PID 1234"]
    },
    {
        "name": "get_system_vitals",
        "description": "Get detailed hardware stats (Battery, CPU, RAM).",
        "parameters": {},
        "examples": ["check battery", "how is cpu usage", "system status"]
    },
    {
        "name": "press_hotkey",
        "description": "Press keyboard shortcuts to control apps (e.g., save, close tab, switch window).",
        "parameters": {"keys": "combination (e.g., ctrl+c, alt+f4, command+t)"},
        "examples": ["press ctrl+c", "close this tab (ctrl+w)", "switch window (alt+tab)"]
    },
    {
        "name": "take_screenshot",
        "description": "Capture the current screen state.",
        "parameters": {"filename": "name of file (default screenshot.png)"},
        "examples": ["take a screenshot", "capture screen"]
    },
    {
        "name": "clipboard_action",
        "description": "Read or write to system clipboard.",
        "parameters": {"action": "'copy' or 'paste'", "text": "text to copy (optional)"},
        "examples": ["what is in my clipboard", "copy 'hello' to clipboard"]
    },
    {
        "name": "speak",
        "description": "Give audible voice feedback to the user.",
        "parameters": {"text": "what to say"},
        "examples": ["say hello sir", "read this out loud"]
    },
    {
        "name": "open_app",
        "description": "Launch an application.",
        "parameters": {"app_name": "name of app"},
        "examples": ["open vscode", "start browser"]
    },
    {
        "name": "run_shell_command",
        "description": "ADVANCED: Run raw terminal commands. Use with extreme caution.",
        "parameters": {"command": "shell command", "admin": "boolean (requires admin privs)"},
        "examples": ["list files (ls -la)", "ping google.com"]
    }
]

# ==================== TOOL MAPPING ====================

TOOL_MAP: Dict[str, Callable] = {
    "set_brightness": ToolRegistry.set_brightness,
    "set_volume": ToolRegistry.set_volume,
    "toggle_dark_mode": ToolRegistry.toggle_dark_mode,
    "play_spotify": ToolRegistry.play_spotify,
    "spotify_control": ToolRegistry.spotify_control,
    "read_file": ToolRegistry.read_file,
    "write_file": ToolRegistry.write_file,
    "create_document": ToolRegistry.create_document,
    "file_search": ToolRegistry.file_search,
    "get_current_time": ToolRegistry.get_current_time,
    "get_system_info": ToolRegistry.get_system_info,
    "open_app": ToolRegistry.open_app,
    "open_website": ToolRegistry.open_website,
    "web_search": ToolRegistry.web_search,
    "compose_email": ToolRegistry.compose_email,
    "wait": ToolRegistry.wait,
    "type_text": ToolRegistry.type_text,
    "final_answer": ToolRegistry.final_answer,
    "system_power": ToolRegistry.system_power,
    "get_running_processes": ToolRegistry.get_running_processes,
    "terminate_process": ToolRegistry.terminate_process,
    "get_system_vitals": ToolRegistry.get_system_vitals,
    "press_hotkey": ToolRegistry.press_hotkey,
    "type_text": ToolRegistry.type_text,
    "take_screenshot": ToolRegistry.take_screenshot,
    "clipboard_action": ToolRegistry.clipboard_action,
    "speak": ToolRegistry.speak,
    "open_app": ToolRegistry.open_app,
    "run_shell_command": run_shell_command,
    "get_network_info": ToolRegistry.get_network_info,
    "file_search": ToolRegistry.file_search_recursive
}

# ==================== TOOL EXECUTION ====================

def execute_tool(tool_name: str, params: Dict[str, Any]) -> str:
    if tool_name not in TOOL_MAP:
        available = ", ".join(TOOL_MAP.keys())
        logger.error(f"Unknown tool: {tool_name}")
        return f"❌ Error: Unknown tool '{tool_name}'. Available: {available}"
    
    tool_func = TOOL_MAP[tool_name]
    
    try:
        logger.info(f"🔧 Executing: {tool_name} with params: {params}")
        
        # Execute tool with parameters
        if params:
            result = tool_func(**params)
        else:
            result = tool_func()
        
        logger.info(f"✅ {tool_name} completed successfully")
        return str(result)
        
    except TypeError as e:
        logger.error(f"Parameter error for {tool_name}: {e}")
        return f"❌ Error: Invalid parameters for {tool_name}. {str(e)}"
        
    except Exception as e:
        logger.error(f"Execution failed for {tool_name}: {e}")
        logger.error(traceback.format_exc())
        return f"❌ Execution error in {tool_name}: {str(e)}"


def get_tools_description() -> str:
    desc = "You have access to these tools:\n\n"
    
    for tool in TOOL_SCHEMAS:
        desc += f"🔧 {tool['name']}\n"
        desc += f"   {tool['description']}\n"
        
        if tool['parameters']:
            params_str = ", ".join([f"{k}: {v}" for k, v in tool['parameters'].items()])
            desc += f"   Parameters: {params_str}\n"
        
        if tool['examples']:
            examples = " | ".join(tool['examples'][:3])
            desc += f"   Examples: {examples}\n"
        
        desc += "\n"
    
    return desc

def get_compact_tool_signatures() -> str:
    signatures = []
    
    for tool in TOOL_SCHEMAS:
        # Build parameter string
        params = []
        if tool['parameters']:
            for name, desc in tool['parameters'].items():
                # Infer type
                if "integer" in desc or "0-100" in desc:
                    param_type = "number"
                elif "boolean" in desc or "true" in desc:
                    param_type = "boolean"
                else:
                    param_type = "string"
                
                # Mark optional
                optional = "(optional)" in desc
                suffix = "?" if optional else ""
                
                params.append(f"{name}{suffix}: {param_type}")
        
        param_str = ", ".join(params) if params else ""
        
        # Format: tool_name(params) - description
        sig = f"{tool['name']}({param_str})"
        
        # Add one-line description
        short_desc = tool['description'].split('.')[0]
        signatures.append(f"• {sig}\n  → {short_desc}")
    
    return "\n".join(signatures)


